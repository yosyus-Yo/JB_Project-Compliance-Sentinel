"""Multimodal input — 파일에서 텍스트 추출 (offline-first, 비용 0).

지원 형식:
  - PDF (.pdf)                 → pdfplumber
  - DOCX (.docx)               → python-docx
  - XLSX (.xlsx)               → openpyxl
  - RTF (.rtf)                 → striprtf
  - HTML (.html, .htm)         → BeautifulSoup
  - HWPX (.hwpx)               → stdlib zipfile + xml.etree (외부 lib 0)
  - 이미지 (.png/.jpg/.jpeg/.tiff/.bmp) → pytesseract OCR (tesseract OS 설치 필요)
  - 텍스트 (.txt, .md, .json, .csv) → 기본 텍스트 read

설계 원칙:
  - **lazy import**: 각 형식 의존성은 호출 시점에 import → 미설치여도 다른 형식 동작
  - **size guard**: MAX_BYTES 초과 시 차단 (DoS 방어)
  - **encoding 안전**: utf-8 → cp949(한국어 windows) fallback
  - **dispatcher**: suffix 기반 자동 라우팅 (extract_text_from_path / from_bytes)

설치:
  pip install -e ".[multimodal]"
  # 이미지 OCR 추가: brew install tesseract tesseract-lang  (또는 apt-get install)
"""
from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

MAX_BYTES = 20 * 1024 * 1024  # 20 MB — DoS 방어
SUPPORTED_EXTENSIONS = {
    "txt", "md", "json", "csv",  # plain text (이미 streamlit_app에서 지원)
    "pdf", "docx", "xlsx", "rtf",
    "html", "htm", "hwpx",
    "png", "jpg", "jpeg", "tiff", "bmp",
}


@dataclass(frozen=True)
class ExtractedDocument:
    """파일에서 추출된 텍스트 + 메타데이터."""
    text: str
    source_filename: str
    suffix: str
    extractor: str  # pdfplumber/python-docx/...
    page_count: int = 0
    char_count: int = 0
    warnings: list[str] = field(default_factory=list)


class MultimodalExtractError(Exception):
    """추출 실패 — 의존성 미설치/포맷 오류/크기 초과 등."""


# ──────────────────────────────────────────────────────────────
# 형식별 추출 함수 (각각 lazy import)
# ──────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes, filename: str) -> ExtractedDocument:
    try:
        import pdfplumber  # type: ignore
    except ImportError as e:
        raise MultimodalExtractError(
            "pdfplumber 미설치. `pip install -e \".[multimodal]\"` 또는 `pip install pdfplumber`"
        ) from e
    pages_text: list[str] = []
    warnings: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            try:
                text = page.extract_text() or ""
                pages_text.append(text)
            except Exception as exc:  # 페이지별 추출 실패는 silent skip
                warnings.append(f"page {page.page_number} skip: {exc}")
    full = "\n\n".join(pages_text)
    return ExtractedDocument(
        text=full, source_filename=filename, suffix="pdf",
        extractor="pdfplumber",
        page_count=len(pages_text),
        char_count=len(full),
        warnings=warnings,
    )


def _extract_docx(data: bytes, filename: str) -> ExtractedDocument:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise MultimodalExtractError("python-docx 미설치. `pip install python-docx`") from e
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # table 안의 텍스트도 포함
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    full = "\n".join(paragraphs)
    return ExtractedDocument(
        text=full, source_filename=filename, suffix="docx",
        extractor="python-docx",
        page_count=0,
        char_count=len(full),
    )


def _extract_xlsx(data: bytes, filename: str) -> ExtractedDocument:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError as e:
        raise MultimodalExtractError("openpyxl 미설치. `pip install openpyxl`") from e
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    full = "\n".join(parts)
    return ExtractedDocument(
        text=full, source_filename=filename, suffix="xlsx",
        extractor="openpyxl",
        page_count=len(wb.worksheets),
        char_count=len(full),
    )


def _extract_rtf(data: bytes, filename: str) -> ExtractedDocument:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
    except ImportError as e:
        raise MultimodalExtractError("striprtf 미설치. `pip install striprtf`") from e
    # RTF는 ASCII 기반이므로 utf-8 → cp949 → latin-1 순으로 시도
    text = _decode_safely(data)
    extracted = rtf_to_text(text)
    return ExtractedDocument(
        text=extracted, source_filename=filename, suffix="rtf",
        extractor="striprtf",
        char_count=len(extracted),
    )


def _extract_html(data: bytes, filename: str) -> ExtractedDocument:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError as e:
        raise MultimodalExtractError("beautifulsoup4 미설치. `pip install beautifulsoup4`") from e
    text = _decode_safely(data)
    soup = BeautifulSoup(text, "html.parser")
    # script/style 제거
    for tag in soup(["script", "style"]):
        tag.decompose()
    extracted = soup.get_text(separator="\n", strip=True)
    return ExtractedDocument(
        text=extracted, source_filename=filename, suffix="html",
        extractor="beautifulsoup4",
        char_count=len(extracted),
    )


def _extract_hwpx(data: bytes, filename: str) -> ExtractedDocument:
    """HWPX (한국 공공기관 표준) — zip 컨테이너 + XML 본문.

    스펙: 한컴 HWPX 5.0 공식 (zip + Contents/section*.xml + 기타 메타)
    외부 lib 불필요 — Python stdlib zipfile + xml.etree만 사용.
    """
    parts: list[str] = []
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            # section0.xml, section1.xml, ... 순서대로 읽음
            section_names = sorted(
                n for n in zf.namelist()
                if n.startswith("Contents/section") and n.endswith(".xml")
            )
            if not section_names:
                # 일부 hwpx는 Preview/PrvText.txt에 미리보기 텍스트 보유 (저품질 fallback)
                preview = "Preview/PrvText.txt"
                if preview in zf.namelist():
                    text = zf.read(preview).decode("utf-8", errors="replace")
                    warnings.append("body xml 없음 — Preview/PrvText.txt fallback")
                    return ExtractedDocument(
                        text=text, source_filename=filename, suffix="hwpx",
                        extractor="hwpx-stdlib-preview",
                        char_count=len(text),
                        warnings=warnings,
                    )
                raise MultimodalExtractError("HWPX 본문 section xml 없음 + Preview fallback 부재")

            for section in section_names:
                xml_bytes = zf.read(section)
                try:
                    root = ET.fromstring(xml_bytes)
                except ET.ParseError as exc:
                    warnings.append(f"{section} parse fail: {exc}")
                    continue
                # HWPX 5.0: 본문 텍스트는 hp:t 엘리먼트 안 (namespace 있음)
                # 모든 <t> 텍스트 수집 (namespace-agnostic)
                for elem in root.iter():
                    tag = elem.tag.split("}", 1)[-1]  # namespace 제거
                    if tag == "t" and elem.text:
                        parts.append(elem.text)
                parts.append("")  # section 간 구분
    except zipfile.BadZipFile as e:
        raise MultimodalExtractError(f"HWPX zip 컨테이너 손상: {e}") from e

    full = "\n".join(parts).strip()
    return ExtractedDocument(
        text=full, source_filename=filename, suffix="hwpx",
        extractor="hwpx-stdlib",
        char_count=len(full),
        warnings=warnings,
    )


def _extract_image(data: bytes, filename: str, suffix: str) -> ExtractedDocument:
    """이미지 OCR — pytesseract + Pillow.

    tesseract 바이너리 OS 설치 필요 (`brew install tesseract tesseract-lang`).
    한국어/영어 OCR.
    """
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError as e:
        raise MultimodalExtractError(
            "pytesseract/Pillow 미설치. `pip install pytesseract Pillow` + `brew install tesseract tesseract-lang`"
        ) from e
    try:
        image = Image.open(io.BytesIO(data))
    except Exception as e:
        raise MultimodalExtractError(f"이미지 로드 실패: {e}") from e
    try:
        # 한국어+영어 동시 OCR (kor+eng 언어팩 필요)
        text = pytesseract.image_to_string(image, lang="kor+eng")
    except pytesseract.TesseractNotFoundError as e:
        raise MultimodalExtractError(
            "tesseract 바이너리 미설치. `brew install tesseract tesseract-lang` (macOS) "
            "또는 `apt-get install tesseract-ocr tesseract-ocr-kor` (Linux)"
        ) from e
    except Exception as e:
        # 언어팩 없으면 eng만으로 retry
        try:
            text = pytesseract.image_to_string(image, lang="eng")
        except Exception as e2:
            raise MultimodalExtractError(f"OCR 실패: {e2}") from e2
    return ExtractedDocument(
        text=text, source_filename=filename, suffix=suffix,
        extractor="pytesseract",
        char_count=len(text),
    )


def _extract_plain_text(data: bytes, filename: str, suffix: str) -> ExtractedDocument:
    """txt/md/json/csv — 기본 텍스트 디코딩."""
    text = _decode_safely(data)
    return ExtractedDocument(
        text=text, source_filename=filename, suffix=suffix,
        extractor="plain-text",
        char_count=len(text),
    )


# ──────────────────────────────────────────────────────────────
# 공통 유틸
# ──────────────────────────────────────────────────────────────

def _decode_safely(data: bytes) -> str:
    """UTF-8 → CP949 → Latin-1 순서로 디코딩 시도. 한국어 안전."""
    for encoding in ("utf-8", "utf-8-sig", "cp949", "euc-kr", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # 모두 실패 시 errors="replace"
    return data.decode("utf-8", errors="replace")


# ──────────────────────────────────────────────────────────────
# Dispatcher
# ──────────────────────────────────────────────────────────────

_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "rtf": _extract_rtf,
    "html": _extract_html,
    "htm": _extract_html,
    "hwpx": _extract_hwpx,
    "png": lambda d, f: _extract_image(d, f, "png"),
    "jpg": lambda d, f: _extract_image(d, f, "jpg"),
    "jpeg": lambda d, f: _extract_image(d, f, "jpeg"),
    "tiff": lambda d, f: _extract_image(d, f, "tiff"),
    "bmp": lambda d, f: _extract_image(d, f, "bmp"),
    "txt": lambda d, f: _extract_plain_text(d, f, "txt"),
    "md": lambda d, f: _extract_plain_text(d, f, "md"),
    "json": lambda d, f: _extract_plain_text(d, f, "json"),
    "csv": lambda d, f: _extract_plain_text(d, f, "csv"),
}


def extract_text_from_bytes(data: bytes, filename: str) -> ExtractedDocument:
    """파일 바이트 + 파일명에서 텍스트 추출.

    Streamlit UploadedFile 핸들링용 entry point.
    """
    if len(data) > MAX_BYTES:
        raise MultimodalExtractError(
            f"파일 크기 초과: {len(data) / 1024 / 1024:.1f}MB > {MAX_BYTES / 1024 / 1024:.0f}MB"
        )
    suffix = Path(filename).suffix.lower().lstrip(".")
    if not suffix:
        raise MultimodalExtractError(f"확장자 없음: {filename}")
    extractor = _EXTRACTORS.get(suffix)
    if not extractor:
        raise MultimodalExtractError(
            f"미지원 형식: .{suffix} (지원: {', '.join(sorted(SUPPORTED_EXTENSIONS))})"
        )
    return extractor(data, filename)


def extract_text_from_path(path: Path | str) -> ExtractedDocument:
    """파일 경로에서 텍스트 추출 — CLI/배치 처리용."""
    p = Path(path)
    if not p.exists():
        raise MultimodalExtractError(f"파일 없음: {p}")
    data = p.read_bytes()
    return extract_text_from_bytes(data, p.name)


def is_supported(filename: str) -> bool:
    """파일명 확장자가 지원 형식인지 확인."""
    suffix = Path(filename).suffix.lower().lstrip(".")
    return suffix in SUPPORTED_EXTENSIONS
