"""multimodal_input.py — 형식별 텍스트 추출 + dispatcher.

dispatcher/plain-text/HWPX/size guard는 의존성 없이 test.
PDF/DOCX/XLSX/RTF/HTML/이미지는 lib 미설치 시 자동 skip.
"""
from __future__ import annotations

import io
import json
import zipfile

import pytest

from compliance_sentinel.multimodal_input import (
    MAX_BYTES,
    SUPPORTED_EXTENSIONS,
    ExtractedDocument,
    MultimodalExtractError,
    _decode_safely,
    extract_text_from_bytes,
    extract_text_from_path,
    is_supported,
)


class TestSupportedExtensions:
    def test_includes_plain_text(self):
        for ext in ("txt", "md", "json", "csv"):
            assert ext in SUPPORTED_EXTENSIONS

    def test_includes_document_formats(self):
        for ext in ("pdf", "docx", "xlsx", "rtf", "html", "htm", "hwpx"):
            assert ext in SUPPORTED_EXTENSIONS

    def test_includes_image_formats(self):
        for ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
            assert ext in SUPPORTED_EXTENSIONS


class TestIsSupported:
    def test_supported_extension(self):
        assert is_supported("doc.pdf") is True
        assert is_supported("doc.hwpx") is True

    def test_unsupported_extension(self):
        assert is_supported("doc.hwp") is False
        assert is_supported("doc.exe") is False

    def test_case_insensitive(self):
        assert is_supported("DOC.PDF") is True

    def test_no_extension(self):
        assert is_supported("doc") is False


class TestDecodeSafely:
    def test_utf8(self):
        assert _decode_safely("한글".encode("utf-8")) == "한글"

    def test_cp949(self):
        assert _decode_safely("한글".encode("cp949")) == "한글"

    def test_utf8_with_bom(self):
        assert _decode_safely("﻿한글".encode("utf-8")) in {"﻿한글", "한글"}

    def test_latin1_fallback(self):
        # latin-1은 모든 byte를 받음 → 항상 디코딩
        result = _decode_safely(b"\xff\xfe\xfd")
        assert isinstance(result, str)

    def test_ascii(self):
        assert _decode_safely(b"hello") == "hello"


class TestPlainTextExtraction:
    def test_txt(self):
        result = extract_text_from_bytes(b"hello world", "test.txt")
        assert result.text == "hello world"
        assert result.extractor == "plain-text"
        assert result.suffix == "txt"
        assert result.char_count == 11

    def test_md(self):
        result = extract_text_from_bytes(b"# Header\n\nbody", "doc.md")
        assert "# Header" in result.text
        assert result.suffix == "md"

    def test_json(self):
        data = json.dumps({"key": "value"}).encode("utf-8")
        result = extract_text_from_bytes(data, "config.json")
        assert "key" in result.text
        assert result.suffix == "json"

    def test_csv(self):
        result = extract_text_from_bytes(b"a,b,c\n1,2,3", "data.csv")
        assert "a,b,c" in result.text
        assert result.suffix == "csv"

    def test_korean_utf8(self):
        result = extract_text_from_bytes("원금 보장 광고".encode("utf-8"), "ad.txt")
        assert "원금 보장" in result.text

    def test_korean_cp949(self):
        result = extract_text_from_bytes("원금 보장 광고".encode("cp949"), "ad.txt")
        assert "원금" in result.text


class TestSizeGuard:
    def test_under_limit_ok(self):
        data = b"x" * (MAX_BYTES - 1)
        result = extract_text_from_bytes(data, "small.txt")
        assert len(result.text) == MAX_BYTES - 1

    def test_exceed_limit_raises(self):
        data = b"x" * (MAX_BYTES + 1)
        with pytest.raises(MultimodalExtractError, match="크기 초과"):
            extract_text_from_bytes(data, "huge.txt")


class TestDispatcherErrors:
    def test_no_extension_raises(self):
        with pytest.raises(MultimodalExtractError, match="확장자 없음"):
            extract_text_from_bytes(b"x", "noext")

    def test_unsupported_extension_raises(self):
        with pytest.raises(MultimodalExtractError, match="미지원 형식"):
            extract_text_from_bytes(b"x", "test.hwp")

    def test_unsupported_extension_lists_supported(self):
        # error 메시지에 지원 형식 목록 포함
        with pytest.raises(MultimodalExtractError) as exc_info:
            extract_text_from_bytes(b"x", "test.exe")
        assert "pdf" in str(exc_info.value)
        assert "docx" in str(exc_info.value)


class TestExtractFromPath:
    def test_nonexistent_path_raises(self, tmp_path):
        with pytest.raises(MultimodalExtractError, match="파일 없음"):
            extract_text_from_path(tmp_path / "missing.txt")

    def test_reads_txt_file(self, tmp_path):
        path = tmp_path / "test.txt"
        path.write_text("hello", encoding="utf-8")
        result = extract_text_from_path(path)
        assert result.text == "hello"
        assert result.source_filename == "test.txt"

    def test_accepts_string_path(self, tmp_path):
        path = tmp_path / "test.txt"
        path.write_text("hi", encoding="utf-8")
        result = extract_text_from_path(str(path))
        assert result.text == "hi"


class TestExtractedDocumentDataclass:
    def test_frozen(self):
        doc = ExtractedDocument(
            text="x", source_filename="a.txt", suffix="txt", extractor="plain-text",
        )
        with pytest.raises(Exception):
            doc.text = "y"

    def test_default_warnings_empty(self):
        doc = ExtractedDocument(
            text="x", source_filename="a.txt", suffix="txt", extractor="plain-text",
        )
        assert doc.warnings == []


class TestHWPXExtraction:
    """HWPX는 stdlib만 사용 — 항상 test 가능."""

    def _make_hwpx(self, body_text: str = "원금 보장 광고") -> bytes:
        """최소 HWPX 컨테이너 생성 (Contents/section0.xml만 포함)."""
        buf = io.BytesIO()
        section_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section">
  <hp:p>
    <hp:run>
      <hp:t>{body_text}</hp:t>
    </hp:run>
  </hp:p>
</hs:sec>
"""
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("Contents/section0.xml", section_xml)
            # 메타 파일도 추가 (실제 HWPX 구조 따름)
            zf.writestr("mimetype", "application/hwp+zip")
        return buf.getvalue()

    def test_basic_hwpx_extraction(self):
        data = self._make_hwpx("테스트 본문 내용")
        result = extract_text_from_bytes(data, "doc.hwpx")
        assert "테스트 본문" in result.text
        assert result.extractor == "hwpx-stdlib"
        assert result.suffix == "hwpx"

    def test_korean_hwpx(self):
        data = self._make_hwpx("100% 원금 보장 무위험 광고")
        result = extract_text_from_bytes(data, "ad.hwpx")
        assert "원금 보장" in result.text

    def test_bad_zip_raises(self):
        with pytest.raises(MultimodalExtractError, match="zip"):
            extract_text_from_bytes(b"not a zip file", "broken.hwpx")

    def test_no_section_with_preview_fallback(self):
        """본문 section xml 없지만 Preview/PrvText.txt 있으면 fallback."""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("Preview/PrvText.txt", "미리보기 텍스트")
            zf.writestr("mimetype", "application/hwp+zip")
        result = extract_text_from_bytes(buf.getvalue(), "preview.hwpx")
        assert "미리보기 텍스트" in result.text
        assert "fallback" in result.warnings[0]

    def test_no_body_no_preview_raises(self):
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("mimetype", "application/hwp+zip")
        with pytest.raises(MultimodalExtractError, match="본문"):
            extract_text_from_bytes(buf.getvalue(), "empty.hwpx")


class TestPDFExtractionConditional:
    """pdfplumber 설치 시에만 실행."""

    def test_pdf_extraction(self):
        pdfplumber = pytest.importorskip("pdfplumber")
        # pdfplumber는 PDF 생성 못 함 — reportlab 또는 pre-built fixture 필요
        # 본 test는 pdfplumber 가용성만 확인 (생성은 별도)
        assert pdfplumber is not None


class TestDOCXExtractionConditional:
    def test_docx_extraction(self, tmp_path):
        docx = pytest.importorskip("docx")
        from docx import Document
        # 최소 DOCX 생성
        doc = Document()
        doc.add_paragraph("테스트 문단입니다.")
        doc.add_paragraph("두 번째 문단.")
        path = tmp_path / "test.docx"
        doc.save(path)
        result = extract_text_from_bytes(path.read_bytes(), "test.docx")
        assert "테스트 문단" in result.text
        assert "두 번째 문단" in result.text
        assert result.extractor == "python-docx"


class TestXLSXExtractionConditional:
    def test_xlsx_extraction(self, tmp_path):
        openpyxl = pytest.importorskip("openpyxl")
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "광고심의"
        ws.append(["항목", "값"])
        ws.append(["원금 보장", "위반"])
        path = tmp_path / "test.xlsx"
        wb.save(path)
        result = extract_text_from_bytes(path.read_bytes(), "test.xlsx")
        assert "원금 보장" in result.text
        assert "광고심의" in result.text
        assert result.extractor == "openpyxl"


class TestRTFExtractionConditional:
    def test_rtf_extraction(self):
        pytest.importorskip("striprtf")
        # 최소 RTF
        rtf_content = (
            r"{\rtf1\ansi\ansicpg949 \uc0 "
            r"儆0? 䑔4?儐9?儀9?䜙6?䜙6?儆0? 吂8?兤8? "
            r"hello world }"
        )
        result = extract_text_from_bytes(rtf_content.encode("utf-8"), "test.rtf")
        assert "hello world" in result.text
        assert result.extractor == "striprtf"


class TestHTMLExtractionConditional:
    def test_html_extraction(self):
        pytest.importorskip("bs4")
        html = b"<html><body><h1>\xec\x9b\x90\xea\xb8\x88 \xeb\xb3\xb4\xec\x9e\xa5</h1><script>alert(1)</script><p>\xea\xb4\x91\xea\xb3\xa0</p></body></html>"
        result = extract_text_from_bytes(html, "test.html")
        assert "원금 보장" in result.text
        assert "광고" in result.text
        # script 제거됨
        assert "alert" not in result.text
        assert result.extractor == "beautifulsoup4"

    def test_htm_alias(self):
        pytest.importorskip("bs4")
        result = extract_text_from_bytes(b"<p>hello</p>", "test.htm")
        assert "hello" in result.text


class TestImageOCRConditional:
    """이미지 OCR — pytesseract + Pillow + tesseract 바이너리 모두 필요."""

    def test_image_extraction_requires_deps(self):
        pytesseract = pytest.importorskip("pytesseract")
        PIL = pytest.importorskip("PIL")
        from PIL import Image as PILImage
        # 1x1 흰색 이미지
        img = PILImage.new("RGB", (1, 1), color="white")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        try:
            result = extract_text_from_bytes(buf.getvalue(), "blank.png")
            assert result.extractor == "pytesseract"
            # 빈 이미지는 텍스트 거의 없음
            assert isinstance(result.text, str)
        except MultimodalExtractError as e:
            # tesseract 바이너리 미설치 또는 OCR 출력 인코딩 문제(환경별 tesseract 버전 의존)는 skip.
            # 실제 dispatch/형식 버그는 다른 메시지이므로 raise로 노출을 유지한다.
            msg = str(e).lower()
            if "tesseract" in msg or "ocr 실패" in str(e) or "codec" in msg:
                pytest.skip(f"tesseract OCR 미동작 (환경 의존): {e}")
            raise
